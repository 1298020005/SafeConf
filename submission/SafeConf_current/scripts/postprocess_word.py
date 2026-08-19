#!/usr/bin/env python3
"""Apply journal-facing Word styles after Pandoc conversion."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HERE = Path(__file__).resolve().parents[1]


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = paragraph._p.get_or_add_pPr()
    suppress_line_number = OxmlElement("w:suppressLineNumbers")
    p_pr.append(suppress_line_number)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    set_run_font(run, "Arial", 9)


def add_continuous_line_numbers(section) -> None:
    sect_pr = section._sectPr
    for node in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(node)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:distance"), "360")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def style_research_document(path: Path, *, double_spaced: bool) -> None:
    document = Document(path)

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_page_number(section.footer.paragraphs[0])
        if double_spaced:
            add_continuous_line_numbers(section)

    for style_name in ("Normal", "Body Text", "First Paragraph"):
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 2 if double_spaced else 1.15
        style.paragraph_format.space_after = Pt(0)

    heading_sizes = {"Title": 16, "Heading 1": 14, "Heading 2": 12, "Heading 3": 11}
    for style_name, size in heading_sizes.items():
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.line_spacing = 2 if double_spaced else 1.15

    if "Caption" in document.styles:
        style = document.styles["Caption"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing = 2 if double_spaced else 1.15

    for paragraph in document.paragraphs:
        if double_spaced:
            paragraph.paragraph_format.line_spacing = 2
            paragraph.paragraph_format.space_after = Pt(0)
        if paragraph.text.strip().startswith("Table "):
            paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            if paragraph.style.name.startswith("Heading"):
                set_run_font(run, "Arial")
            else:
                set_run_font(run, "Times New Roman")

    for table in document.tables:
        if "Table Grid" in document.styles:
            table.style = "Table Grid"
        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
            for cell in row.cells:
                # BMC requests tables without colour or shading.
                tc_pr = cell._tc.get_or_add_tcPr()
                for shading in tc_pr.findall(qn("w:shd")):
                    tc_pr.remove(shading)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_run_font(run, "Times New Roman", 8.5)
                        if row_index == 0:
                            run.font.bold = True

    document.save(path)


def main() -> None:
    style_research_document(HERE / "SafeConf_manuscript.docx", double_spaced=True)
    style_research_document(HERE / "SafeConf_supplement.docx", double_spaced=True)
    style_research_document(
        HERE / "Cover_letter_BMC_Bioinformatics.docx", double_spaced=False
    )
    print("Applied Word margins, fonts, page numbers, and continuous line numbering")


if __name__ == "__main__":
    main()
